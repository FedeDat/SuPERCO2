#!/usr/bin/env python
# coding: utf-8

# In[9]:


#!/usr/bin/env python
# coding: utf-8

import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
import shutil
from scipy.signal import savgol_filter
from scipy.signal import find_peaks
from scipy.optimize import curve_fit
import math
from numpy import trapz
import requests

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)

# Functions

def cm_to_inch(cm):
    inch=float(cm/2.54)
    return inch

def AgCl_to_RHE(V_AgCl,pH,ref):
    if ref == "Ag/AgCl":
        V_RHE=+0.197+V_AgCl+0.059*pH
        return V_RHE


# In[10]:


# Ask the user for inputs
cat = input("Enter the catalyst name (e.g., CuO):\n").strip()
date = input("Enter the date in YYYY-MM-DD format (e.g., 2025-03-05).\nAdd _n if there are repeated experiments:\n").strip()


# In[11]:


# Initialize input
url_OCV = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/OCV/OCV.csv"
url_CV_CO2 = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/CV/CV-CO2.csv"
url_CV_N2 = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/CV/CV-N2.csv"
url_CA = f"https://api.github.com/repos/FedeDat/SuPERCO2/contents/data/{cat}/electrochemistry/{date}/CA"


# # Open circuit potential

# In[13]:


# Read the content of the file

data_OCV = pd.read_csv(url_OCV, header=None, names=['time', 'V'])

pH_OCV=float(data_OCV['V'][0])
ref_OCV=str(data_OCV['V'][1])

data_OCV=data_OCV[4:]

data_OCV=data_OCV.astype(float)

data_OCV['V']=AgCl_to_RHE(data_OCV['V'],6.8,ref_OCV)

plt.figure(figsize=(cm_to_inch(12), cm_to_inch(8)))

plt.plot(data_OCV['time'],data_OCV['V'], )

plt.title('Open circuit potential')

plt.xlabel(r"$t$ (min)")
plt.ylabel(r"$V$ (vs RHE)")

plt.ylim(min(np.min(data_OCV['V']),0.5*np.mean(data_OCV['V'])),max(np.max(data_OCV['V']),1.5*np.mean(data_OCV['V'])))

plt.savefig('OCV.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# # Comparison of cyclic voltammetries in CO<sub>2</sub> and N<sub>2</sub>

# In[14]:


data_CV_CO2 = pd.read_csv(url_CV_CO2, header=None, names=['V', 'j'])

pH_CV_CO2=float(data_CV_CO2['j'][0])
ref_CV_CO2=str(data_CV_CO2['j'][1])

data_CV_CO2=data_CV_CO2[4:]

data_CV_CO2=data_CV_CO2.astype(float)

data_CV_CO2['V']=AgCl_to_RHE(data_CV_CO2['V'],6.8,ref_CV_CO2)

data_CV_N2 = pd.read_csv(url_CV_N2, header=None, names=['V', 'j'])

pH_CV_N2=float(data_CV_N2['j'][0])
ref_CV_N2=str(data_CV_N2['j'][1])

data_CV_N2=data_CV_N2[4:]

data_CV_N2=data_CV_N2.astype(float)

data_CV_N2['V']=AgCl_to_RHE(data_CV_N2['V'],6.8,ref_CV_N2)

plt.figure(figsize=(cm_to_inch(12), cm_to_inch(12)))

plt.plot(data_CV_CO2['V'],data_CV_CO2['j'], label=r"CO$_2$ saturated",color='red')
plt.plot(data_CV_N2['V'],data_CV_N2['j'], label=r"N$_2$ saturated",color='gray')

plt.title('Cyclic voltammetry')
plt.legend(loc='best')

plt.xlabel(r"$V$ (vs RHE)")

plt.ylabel(r"$j$ (mA/cm$^2$)")

plt.savefig('CV.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# # Calculation of CO<sub>2</sub> reduction activity and selectivity

# In[15]:


# Analyse Faradaic Efficiency

# Get the file listing
response = requests.get(url_CA)
files = response.json()

metadata_EC = pd.DataFrame(columns=['pH', 'ref', 'V'])
anl_EC = pd.DataFrame(columns=['V_RHE','Flow Rate','Mode','Average j', 'Q (mC)'])

data = pd.DataFrame()

n1=0

plt.figure(figsize=(cm_to_inch(18), cm_to_inch(12)))

for file in files:
    if file["name"].startswith("CA"):
        raw_url_CA = file["download_url"]  # direct raw link
        data_EC_raw = pd.read_csv(raw_url_CA, header=None, names=['t', 'j'])
        
        metadata_EC.loc[n1,'pH']=float(data_EC_raw['j'][0])
        metadata_EC.loc[n1,'ref']=str(data_EC_raw['j'][1])
        anl_EC.loc[n1,'V_RHE']=round(AgCl_to_RHE(float(data_EC_raw['j'][2]),float(data_EC_raw['j'][0]),str(data_EC_raw['j'][1])),2)
        anl_EC.loc[n1,'Flow Rate']=int(data_EC_raw['j'][3])
        anl_EC.loc[n1,'Mode']=str(data_EC_raw['j'][4])

        data_EC_raw = data_EC_raw[80:]; data_EC_raw=data_EC_raw.astype(float)
        anl_EC.loc[n1,'Average j']=round(np.mean(data_EC_raw['j'][data_EC_raw['j']<0]),1)
        
        if anl_EC.loc[n1,'Mode']!= "DPA":
            plt.plot(data_EC_raw['t']/60,data_EC_raw['j'], label=r"$E=$"+str(anl_EC.loc[n1,'V_RHE'])+" V$_{RHE}$, "+str(anl_EC.loc[n1,'Flow Rate'])+" NmL/min")
        
        if float(data_EC_raw['t'][-1:]) < 1799:
            print(os.path.basename(os.getcwd()),filename[0:3]+" ended before the allocated time.")
        
        anl_EC.loc[n1,'Q (mC)']=trapz(-data_EC_raw['j'][data_EC_raw['j']<0],data_EC_raw['t'][data_EC_raw['j']<0])
        
        n1=n1+1   
        
plt.title('Activity')
plt.legend(loc='best')

plt.xlabel(r"$t$ (min)")
plt.ylabel(r"$j$ (mA/cm$^2$)")

plt.gca().invert_yaxis()
plt.xlim(0,30)

plt.savefig('activity.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()

anl_GC_fr = pd.DataFrame(columns=['H2','CO','CH4'])
anl_GC_nmol = pd.DataFrame(columns=['H2','CO','CH4'])
anl_GC_FE = pd.DataFrame(columns=['H2','CO','CH4'])
n_el = {'H2': 2, 'CO': 2, 'CH4': 8} # Number of electrons

p=1 # Pressure given in atm
R=0.082 # Gas constant given in l atm/(mol K) 
C_farad=96485 # Faradaic constant in C/mol e-
temp=293.15 # temperature in K
dt=4+22/60 # min

n2=0

for file in files:
    if file["name"].startswith("GC"):
        raw_url_GC = file["download_url"]  # direct raw link

        data_GC_raw = pd.read_csv(raw_url_GC)
        
        for mol in ['H2','CO','CH4']:
            anl_GC_fr.loc[n2,mol]=trapz(data_GC_raw[mol]/100*anl_EC['Flow Rate'][n2],dx=dt)
            anl_GC_nmol.loc[n2,mol]=p*anl_GC_fr.loc[n2,mol]/1000*1/(R*temp)
            anl_GC_FE.loc[n2,mol]=(100*anl_GC_nmol.loc[n2,mol]*C_farad*n_el[mol])/(anl_EC['Q (mC)'][n2]/1000)
        
        n2=n2+1 

x_labels=[]

for i in np.arange(len(anl_EC['V_RHE'])):
    x_labels.append(str(anl_EC['V_RHE'][i])+" V$_{RHE}$,\n"+str(anl_EC['Flow Rate'][i])+" NmL/min, \n"+str(anl_EC['Mode'][i]))

anl_GC_FE.plot(kind='bar', stacked=True, figsize=(cm_to_inch(18), cm_to_inch(12)))

plt.title('Selectivity')

plt.xticks(ticks=range(len(x_labels)), labels=x_labels, rotation=0)

plt.ylabel('Faradaic Efficiency (%)')

plt.savefig('selectivity.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# In[20]:


# Saving output

with pd.ExcelWriter("FE_EC-data.xlsx", engine="openpyxl") as writer:
    anl_EC.to_excel(writer,
                    sheet_name='Average EC data', 
                    index=False)
    
    anl_GC_fr.to_excel(writer,
                      sheet_name='Flow rate (NmL min-1)',
                      index=False)
    
    anl_GC_nmol.to_excel(writer,
                        sheet_name='Concentration (mol)',
                        index=False)
    
    anl_GC_FE.to_excel(writer,
                      sheet_name='Selectivity (%)',
                      index=False)


# In[21]:


# Create a new Word document
doc = Document()

# Create or get the footer
section = doc.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

# Add "Page " text
#paragraph.text = "Page "

# Add PAGE field
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"  # Word PAGE field
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

paragraph.alignment = 1  # center alignment

#doc.add_heading('Report of analyzed data', level=1)

# Add a paragraph
#doc.add_paragraph("This is a generated figure:")

# Add an image (from file)
doc.add_picture('OCV.png')
doc.add_picture('CV.png')
doc.add_picture('activity.png')
doc.add_picture('selectivity.png')

# Save the document
doc.save('report_EC.docx')

