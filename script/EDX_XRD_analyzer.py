#!/usr/bin/env python
# coding: utf-8

# In[1]:


#!/usr/bin/env python
# coding: utf-8

import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
import shutil
from scipy.signal import savgol_filter
import math

from scipy.signal import find_peaks
from scipy.optimize import curve_fit

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)

# Functions

def cm_to_inch(cm):
    inch=float(cm/2.54)
    return inch

def AgCl_to_RHE(V_AgCl,pH,ref):
    if ref == "Ag/AgCl":
        V_RHE=+0.197+V_AgCl+0.059*pH
        return V_RHE


# In[2]:


# Ask the user for inputs
cat = input("Enter the catalyst name (e.g., CuO):\n").strip()


# # Analysis of EDX file

# In[3]:


# Read the content of the file

# Initialize input
url_EDX = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/characterization/EDX.csv"

data_EDX = pd.read_csv(url_EDX, usecols=lambda column: pd.notnull(column))
data_EDX.dropna(axis=1, how='all', inplace=True)

#Calculate ratio between metal and oxygen

data_EDX[''+str(data_EDX.columns[0])+'/'+str(data_EDX.columns[1])+'']=data_EDX[data_EDX.columns[0]]/data_EDX[data_EDX.columns[1]]

data_EDX.index=list(range(1, len(data_EDX)+1))

ax1 = data_EDX.iloc[:,:-1].plot(kind='bar', stacked=True, figsize=(cm_to_inch(10), cm_to_inch(8)))
plt.legend(loc='upper center')

ax2 = ax1.twinx()  # Create a second y-axis
ax2.plot(data_EDX.index-1, data_EDX[data_EDX.columns[-1]], linestyle='None', marker='o', markersize=8, markerfacecolor='red', markeredgecolor='black')

# Adding titles and labels
ax1.set_xticklabels(list(range(1, len(data_EDX)+1)), rotation=0, ha='right')

ax1.set_title('Atomic content')
ax1.set_xlabel('Site')
ax1.set_ylabel('Content (At.%)')

ax2.yaxis.label.set_color('red')  # Change the y-axis label color
ax2.spines['right'].set_color('red')  # Change the spine color to red
ax2.tick_params(axis='y', colors='red')  # Change the tick color to red

ax2.set_ylabel(''+str(data_EDX.columns[0])+'/'+str(data_EDX.columns[1])+' ratio')

ax1.set_ylim(0, 100)
ax2.set_ylim(0, 1)

# Display the plot
plt.tight_layout()  # Adjust layout

plt.savefig('EDX.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()

data_EDX.to_excel('EDX_analysed.xlsx', index=False)


# # Analysis of XRD file, including smoothing filter and fit analysis

# In[4]:


url_XRD = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/characterization/XRD.csv"

data_XRD = pd.read_csv(url_XRD, header=None, usecols=[0, 1], names=['Angle', 'Intensity'])
    
k_XRD = float(data_XRD[data_XRD.apply(lambda row: row.astype(str).str.contains('K-Alpha1 wavelength')).any(axis=1)]['Intensity'])/10 # X-ray wavelenght

n_XRD = int(data_XRD[data_XRD.apply(lambda row: row.astype(str).str.contains('No. of points')).any(axis=1)]['Intensity'])

data_XRD_sl = data_XRD.iloc[(data_XRD[data_XRD.apply(lambda row: row.astype(str).str.contains('Angle')).any(axis=1)].index[0]+1):]

data_XRD_sl = data_XRD_sl.astype(float)

data_XRD_sl['Intensity'] = data_XRD_sl['Intensity'] - np.min(data_XRD_sl['Intensity'])

data_XRD_sl['Intensity']=data_XRD_sl['Intensity']/np.max(data_XRD_sl['Intensity'])

data_XRD_sl['Intensity_savgol'] = savgol_filter(data_XRD_sl['Intensity'], window_length=10, polyorder=1)

k_XRD = float(data_XRD[data_XRD.apply(lambda row: row.astype(str).str.contains('K-Alpha1 wavelength')).any(axis=1)]['Intensity'])/10 # X-ray wavelenght

n_XRD = int(data_XRD[data_XRD.apply(lambda row: row.astype(str).str.contains('No. of points')).any(axis=1)]['Intensity'])

data_XRD_sl = data_XRD.iloc[(data_XRD[data_XRD.apply(lambda row: row.astype(str).str.contains('Angle')).any(axis=1)].index[0]+1):]

data_XRD_sl = data_XRD_sl.astype(float)

data_XRD_sl['Intensity'] = data_XRD_sl['Intensity'] - np.min(data_XRD_sl['Intensity'])

data_XRD_sl['Intensity']=data_XRD_sl['Intensity']/np.max(data_XRD_sl['Intensity'])

data_XRD_sl['Intensity_savgol'] = savgol_filter(data_XRD_sl['Intensity'], window_length=10, polyorder=1)

plt.figure(figsize=(cm_to_inch(16), cm_to_inch(8)))

plt.plot(data_XRD_sl['Angle'],data_XRD_sl['Intensity'], label='Raw data', alpha=0.5)
plt.plot(data_XRD_sl['Angle'],data_XRD_sl['Intensity_savgol'], label='Savitzky-Golay smoothing')

plt.title('Raw XRD spectra')

plt.xlabel(r"$2{\theta}$ (deg)")

plt.ylabel('Intensity (a.u.)')

plt.savefig('XRD_spectra.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()

# Step 1: Identify peaks
peaks, _ = find_peaks(data_XRD_sl['Intensity_savgol'], height=0.205, distance=6)  # Adjust height threshold as needed

# Step 2: Define a Gaussian function for fitting
def gaussian(x, amp, mean, stddev):
    return amp * np.exp(-(x - mean) ** 2 / (2 * stddev ** 2))

# Step 3: Fit Gaussian to each peak
fitted_params = []

p_saved=0

for peak in peaks:

    # Define the x range for fitting (e.g., around the peak)
    x_fit = data_XRD_sl['Angle'][peak-20:peak+20]
    y_fit = data_XRD_sl['Intensity'][peak-20:peak+20]

    # Fit Gaussian
    initial_guess = [1, data_XRD_sl['Angle'][peak], 0.3]  # Initial guesses for amplitude, mean, and stddev
    try:
        params, _ = curve_fit(gaussian, x_fit, y_fit, p0=initial_guess)
        fitted_params.append(params)
    except RuntimeError:
        print(f"Fit failed for peak at index {peak}")

# Step 4: Plot the data, peaks, and fitted Gaussian

plt.figure(figsize=(cm_to_inch(16), cm_to_inch(8)))

plt.plot(data_XRD_sl['Angle'], data_XRD_sl['Intensity'], color='blue')

# Plot Gaussian fits
for params in fitted_params:
    amp, mean, stddev = params
    plt.plot(data_XRD_sl['Angle'], gaussian(data_XRD_sl['Angle'], amp, mean, stddev))

n_peak=0

scherrer=pd.DataFrame(index=range(len(peaks)),columns=['Peak position (deg)','Crystallite size (nm)'])

for peak in peaks:

    scherrer.iloc[n_peak,0]=fitted_params[n_peak][1]
    scherrer.iloc[n_peak,1]=k_XRD*np.cos(math.radians(fitted_params[n_peak][1]/2))/(2*np.sqrt(2*np.log(2))*math.radians(fitted_params[n_peak][2]))

    plt.text(data_XRD_sl['Angle'][peak], 1.1, f'{fitted_params[n_peak][1]:.0f}'+'\n deg', va='center', fontsize=8, color="red")

    n_peak=n_peak+1

plt.title('XRD spectra with identified peaks')

plt.xlabel(r"$2{\theta}$ (deg)")

plt.ylabel('Intensity (a.u.)')

plt.ylim(0, 1.2)

scherrer.to_excel('crystallites.xlsx', index=False)
plt.savefig('XRD_spectra_fitted.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# In[5]:


# Create a new Word document
doc = Document()

# Create or get the footer
section = doc.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

# Add "Page " text
#paragraph.text = "Page "

# Add PAGE field
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"  # Word PAGE field
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

paragraph.alignment = 1  # center alignment

#doc.add_heading('Report of analyzed data', level=1)

# Add a paragraph
#doc.add_paragraph("This is a generated figure:")

# Add an image (from file)
doc.add_picture('EDX.png')
doc.add_picture('XRD_spectra.png')
doc.add_picture('XRD_spectra_fitted.png')

# Save the document
doc.save('report_CH.docx')

