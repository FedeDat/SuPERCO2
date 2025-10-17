#!/usr/bin/env python
# coding: utf-8

# In[1]:


#!/usr/bin/env python
# coding: utf-8

import re
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import os
import shutil

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scipy.signal import savgol_filter
import math
from numpy import trapz

from scipy.signal import find_peaks
from scipy.optimize import curve_fit

# Functions

def cm_to_inch(cm):
    inch=float(cm/2.54)
    return inch

def ref_to_RHE(V_ref,pH,ref):
    if ref == "Ag/AgCl":
        V_RHE=+0.197+V_ref+0.059*pH
        return V_RHE
    
def interpol(y_2,y_1,x_2,x_1,t):
    b=(y_2-y_1)/(x_2-x_1)
    a=y_1-b*x_1
    y=a+b*t
    return y

title_size = int(14)
general_size = int(12)


# In[2]:


# Ask the user for inputs
cat = input("Enter the catalyst name (e.g., CuO):\n").strip()
date = input("Enter the date in YYYY-MM-DD format (e.g., 2025-03-05).\nAdd _n if there are repeated experiments:\n").strip()


# In[21]:


# Initialize input
url_CP = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/CP1.csv"
url_GC = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/GC1.csv"
url_HPLC = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/HPLC1.csv"
url_dil = f"https://raw.githubusercontent.com/FedeDat/SuPERCO2/main/data/{cat}/electrochemistry/{date}/dilution1.csv"


# In[22]:


# Initalize metadata

metadata_EC_raw = pd.read_csv(url_CP, 
                              header=None, 
                              names=['item','in','fin'],
                              nrows=11, 
                              index_col=False)

metadata_EC = pd.DataFrame(columns=['A', 
                                    'Elect', 
                                    'Conc',
                                    'Load',
                                    'pH_in',
                                    'pH_fin',
                                    'ZIR_in',
                                    'ZIR_fin',
                                    'ref',
                                    'V',
                                    'CO2_fr_in',
                                    'CO2_fr_out',
                                    'El_fr'])

metadata_EC.loc[0,'A']=float(metadata_EC_raw['in'][0])
metadata_EC.loc[0,'Elect']=str(metadata_EC_raw['in'][1])
metadata_EC.loc[0,'Conc']=float(metadata_EC_raw['in'][2])
metadata_EC.loc[0,'Load']=float(metadata_EC_raw['in'][3])
metadata_EC.loc[0,'pH_in']=float(metadata_EC_raw['in'][4])
metadata_EC.loc[0,'pH_fin']=float(metadata_EC_raw['fin'][4])
metadata_EC.loc[0,'ZIR_in']=float(metadata_EC_raw['in'][5])
metadata_EC.loc[0,'ZIR_fin']=float(metadata_EC_raw['fin'][5])
metadata_EC.loc[0,'ref']=str(metadata_EC_raw['in'][6])
metadata_EC.loc[0,'V']=float(metadata_EC_raw['in'][7])
metadata_EC.loc[0,'CO2_fr_in']=float(metadata_EC_raw['in'][8])
metadata_EC.loc[0,'CO2_fr_out']=float(metadata_EC_raw['fin'][8])
metadata_EC.loc[0,'El_fr']=float(metadata_EC_raw['in'][9])


# In[23]:


data_EC_raw = pd.read_csv(url_CP, 
                          header=None, 
                          names=['t','i','Ewe','Ecell'], 
                          skiprows=13)

data_EC = pd.DataFrame(columns=['t','pH','ZIR','j','Ewe','Ecell'])


# In[24]:


data_EC['t']=data_EC_raw['t']/60
data_EC['pH']=interpol(metadata_EC['pH_fin'][0],
                       metadata_EC['pH_in'][0],
                       data_EC['t'].iloc[-1]*60,
                       data_EC['t'][0]*60,
                       data_EC['t']*60)
data_EC['ZIR']=interpol(metadata_EC['ZIR_fin'][0],
                        metadata_EC['ZIR_in'][0],
                        data_EC['t'].iloc[-1]*60,
                        data_EC['t'][0]*60,
                        data_EC['t']*60)
data_EC['j']=data_EC_raw['i']/metadata_EC['A'][0]
data_EC['Ewe']=ref_to_RHE(data_EC_raw['Ewe'],
                          data_EC['pH'],
                          "Ag/AgCl")-data_EC['ZIR']*data_EC_raw['i']/1000
data_EC['Ecell']=data_EC_raw['Ecell']


# where the cathodic potential is converted from Ag/AgCl to RHE reference, including ohmic and pH corrections, in line with Equation 
# 
# \begin{equation}
# E_{\text{RHE}} = E_{\text{Ag/AgCl}} + 0.059\text{pH} - Ri
# \end{equation}

# In[25]:


# Visualizing raw data

plt.figure(figsize=(cm_to_inch(18), cm_to_inch(10)))

plt.plot(data_EC['t'],data_EC['Ewe'], 
         label=r"$j=$"+str(round(data_EC['j'].mean()))+" mA/cm$^2$\n"+str(metadata_EC['Conc'][0])+" M "+str(metadata_EC['Elect'][0])+"\nCO$_2$ flow rate = "+str(int(metadata_EC['CO2_fr_in'][0]))+"mL/min")

plt.title('Working potential vs time',fontsize=title_size)
plt.legend(loc='best',fontsize=general_size)

plt.xlabel(r"$t$ (min)",fontsize=general_size)
plt.ylabel(r"$E$ (V vs RHE)",fontsize=general_size)

plt.savefig('Ewe.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# In[26]:


plt.figure(figsize=(cm_to_inch(18), cm_to_inch(10)))

plt.plot(data_EC['t'],data_EC['Ecell'], 
         label=r"$j=$"+str(round(data_EC['j'].mean()))+" mA/cm$^2$\n"+str(metadata_EC['Conc'][0])+" M "+str(metadata_EC['Elect'][0])+"\nCO$_2$ flow rate = "+str(int(metadata_EC['CO2_fr_in'][0]))+"mL/min")

plt.title('Cell potential vs time',fontsize=title_size)
plt.legend(loc='best',fontsize=general_size)

plt.xlabel(r"$t$ (min)",fontsize=general_size)
plt.ylabel(r"$E$ cell (V)",fontsize=general_size)

plt.savefig('Ecell.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# In[27]:


# Calculating Faradaic Efficiency

data_GC_raw = pd.read_csv(url_GC)
data_HPLC_raw = pd.read_csv(url_HPLC)
data_dil_raw = pd.read_csv(url_dil)

n_GC=len(data_GC_raw)

anl_product_mol_t = pd.DataFrame(0.0,index=range(n_GC),
                                 columns=['t','H2','CO','CH4','C2H4',
                                          'HCOO-','AcO-','MEG','EtOH','PrOH'])
anl_product_FE_t = pd.DataFrame(index=range(n_GC),
                                columns=['t','H2','CO','CH4','C2H4',
                                         'HCOO-','AcO-','MEG','EtOH','PrOH'])

n_el = {'H2': 2, 'CO': 2, 'CH4': 8, 'C2H4': 12,'HCOO-': 2,'AcO-': 8,
        'MEG': 10,'EtOH': 12,'PrOH': 18} # Number of electrons

p=1 # Pressure given in atm
R=0.082 # Gas constant given in l atm/(mol K) 
C_farad=96485 # Faradaic constant in C/mol e-
temp=273.15 # temperature in K


# In[28]:


dt=float(10.5)
v_loop=float(1) # loop volume in mL
t_fill=v_loop/metadata_EC['CO2_fr_out'][0]*60 # filling time in sec
t_k=0

anl_EC_t = pd.DataFrame(index=range(n_GC), 
                        columns=['t','j','dj','Ewe','dEwe','Ecell','dEcell'])

for k in np.arange(0,len(data_GC_raw)):
    data_EC_j=data_EC[(data_EC['t'] >= t_k) & (data_EC['t'] < t_k+dt)]['j']
    data_EC_Ewe=data_EC[(data_EC['t'] >= t_k) & (data_EC['t'] < t_k+dt)]['Ewe']
    data_EC_Ecell=data_EC[(data_EC['t'] >= t_k) & (data_EC['t'] < t_k+dt)]['Ecell']
    
    anl_EC_t['t'][k]=t_k+dt; 
    anl_EC_t['j'][k]=data_EC_j.mean(); anl_EC_t['dj'][k]=data_EC_j.std(ddof=1)/np.sqrt(len(data_EC_j))
    anl_EC_t['Ewe'][k]=data_EC_Ewe.mean(); anl_EC_t['dEwe'][k]=data_EC_Ewe.std(ddof=1)/np.sqrt(len(data_EC_j))
    anl_EC_t['Ecell'][k]=data_EC_Ecell.mean(); anl_EC_t['dEcell'][k]=data_EC_Ecell.std(ddof=1)/np.sqrt(len(data_EC_j))
    
    mol_el_k=-anl_EC_t['j'][k]*1E-3*metadata_EC['A'][0]*t_fill/C_farad
    
    anl_product_mol_t.at[k,'t']=t_k+dt; 
    anl_product_FE_t.at[k,'t']=t_k+dt; 
    
    for mol in ['H2','CO','CH4','C2H4']:
        anl_product_mol_t.at[k,mol]=data_GC_raw.at[k,mol]/1E6*1E-3*p/(R*temp)
        anl_product_FE_t.at[k,mol]=100*anl_product_mol_t.at[k,mol]*n_el[mol]/mol_el_k
    
    t_k=t_k+dt

charge_tot = trapz(-data_EC_raw['i'],data_EC_raw['t'])/(1000)
mol_el_tot = charge_tot/C_farad

for mol in ['HCOO-','AcO-','MEG','EtOH','PrOH']:
    anl_product_mol_t[mol]=data_HPLC_raw[mol][0]/1000*float(data_dil_raw['Dilution factor'])*float(data_dil_raw['Total volume (mL)'][0])/1000
    anl_product_FE_t[mol]=100*anl_product_mol_t[mol][0]*n_el[mol]/mol_el_tot


# The Faradaic efficiency (FE) of a product is defined as the fraction of the total charge that is used to form that product:
# \begin{equation}
#     \text{FE}_i = \frac{n_i \, z_i \, F}{Q_\text{tot}} \times 100\%
#     \label{eq:fe_basic}
# \end{equation}
# 
# where
# - $n_i$ = amount of species $i$ produced (in moles)
# - $z_i$ = number of electrons transferred per molecule of species $i$
# - $F$ = Faraday constant ($96485~\text{C}\times \text{mol}$)
# - $Q_\text{tot}$ = total charge passed during electrolysis (in moles of electrons

# In[29]:


# Selectivity vs time

ax1 = anl_product_FE_t.drop(columns=['t']).plot(kind='bar', 
                                                stacked=True, 
                                                figsize=(cm_to_inch(18), 
                                                         cm_to_inch(10)))
plt.legend(loc='best',ncol=2,fontsize=general_size)

ax2 = ax1.twinx()  # Create a second y-axis

ax2.errorbar(
    x=anl_EC_t['Ewe'].index,      
    y=anl_EC_t['Ewe'],
    yerr=anl_EC_t['dEwe'],
    linestyle='None',
    marker='o',
    markersize=3,
    markerfacecolor='red',
    markeredgecolor='red',
    ecolor='red',
    capsize=3,  # adds little caps on the bars
    elinewidth=1
)

#ax2.plot(anl_EC['Ewe'], linestyle='None', marker='o', markersize=3, markerfacecolor='red', markeredgecolor='black')

# Adding titles and labels

plt.title('Selectivity vs time', fontsize=title_size)

ax1.set_xticklabels(list(anl_product_FE_t['t']), rotation=0, ha='center')

#ax1.set_title('Hydrocarbon partial current density', fontsize=title_size)
ax1.set_xlabel(r'$t$ (min)', fontsize=general_size)
ax1.set_ylabel('Faradaic Efficiency (%)', fontsize=general_size)

ax2.yaxis.label.set_color('red')  # Change the y-axis label color
ax2.spines['right'].set_color('red')  # Change the spine color to red

ax1.tick_params(axis='x', labelsize=general_size)  
ax1.tick_params(axis='y', labelsize=general_size)  
ax2.tick_params(axis='y', colors='red', labelsize=general_size)  # Change the tick color to red

ax2.set_ylabel(r'$E$ (V vs RHE)', fontsize=general_size)

ax1.set_ylim(0, max(100,anl_product_FE_t.drop(columns=['t']).sum(axis=1).max()))
ax2.set_ylim(round(anl_EC_t['Ewe'].min()-0.2,1), 0)

# Display the plot
plt.tight_layout()  # Adjust layout

#plt.show()
plt.savefig('FE_t.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

plt.close()


# In[30]:


anl_product_FE = pd.DataFrame(index=range(2),columns=['H2','CO','CH4','C2H4',
                                                      'HCOO-','AcO-','MEG','EtOH','PrOH'])
anl_EC = pd.DataFrame(index=range(2),columns=['j','Ewe','Ecell'])

anl_product_FE.iloc[0]=anl_product_FE_t.drop(columns=['t']).mean()
anl_product_FE.iloc[1]=1/2*(anl_product_FE_t.drop(columns=['t']).max()-anl_product_FE_t.drop(columns=['t']).min())

anl_EC['j'].iloc[0]=anl_EC_t['j'].mean()
anl_EC['j'].iloc[1]=1/2*(anl_EC_t['j'].max()-anl_EC_t['j'].min())

anl_EC['Ewe'].iloc[0]=anl_EC_t['Ewe'].mean()
anl_EC['Ewe'].iloc[1]=1/2*(anl_EC_t['Ewe'].max()-anl_EC_t['Ewe'].min())

anl_EC['Ecell'].iloc[0]=anl_EC_t['Ecell'].mean()
anl_EC['Ecell'].iloc[1]=1/2*(anl_EC_t['Ecell'].max()-anl_EC_t['Ecell'].min())


# In[31]:


# Selectivity

y_FE=anl_product_FE.iloc[[0]]
dy_FE=anl_product_FE.iloc[[1]]

dy_FE.index = y_FE.index

x_pos = np.arange(len(y_FE))

x_FE_label=r"$j=$"+str(round(data_EC['j'].mean()))+" mA/cm$^2$\n"+str(metadata_EC['Conc'][0])+" M "+str(metadata_EC['Elect'][0])+"\nCO$_2$ flow rate = "+str(int(metadata_EC['CO2_fr_in'][0]))+" mL/min"

bx1 = y_FE.plot(
    kind='bar', 
    width=0.055, 
    stacked=True, 
    yerr=dy_FE, 
        error_kw={'capsize': 4,
                 'elinewidth': 1},
    figsize=(cm_to_inch(18), cm_to_inch(10))
)
plt.legend(loc='best',ncol=1,fontsize=general_size)

bx1.set_xticks([0])  # Position of the bar (for single-row DataFrame it's 0)
bx1.set_xticklabels([x_FE_label], rotation=0, fontsize=general_size)
#bx1.set_xlabel(x_FE_label, rotation=0, fontsize=general_size)

bx1.tick_params(axis='y', labelsize=general_size)
bx1.set_ylabel('Faradaic Efficiency (%)', fontsize=general_size)

bx1.set_ylim(0, max(100,y_FE.sum(axis=1).max()))

bx2 = bx1.twinx()  # Create a second y-axis
bx2.xaxis.set_visible(False) 

bx2.errorbar(
    x=x_pos,      
    y=anl_EC['Ewe'][0],
    yerr=anl_EC['Ewe'][1],
    linestyle='None',
    marker='o',
    markersize=3,
    markerfacecolor='red',
    markeredgecolor='red',
    ecolor='red',
    capsize=3,  # adds little caps on the bars
    elinewidth=1
)

bx2.set_ylim(round(anl_EC['Ewe'][0]-0.2,1), 0)
bx2.yaxis.label.set_color('red')  # Change the y-axis label color
bx2.spines['right'].set_color('red')  # Change the spine color to red

bx2.tick_params(axis='y', colors='red', labelsize=general_size)  # Change the tick color to red
bx2.set_ylabel(r'$E$ (V vs RHE)', fontsize=general_size)

plt.title('Overall selectivity', fontsize=title_size)

plt.savefig('FE.png', format='png', dpi=300, transparent=True, bbox_inches='tight')

# Display the plot
plt.tight_layout()  # Adjust layout

plt.close()


# In[32]:


# Saving output

with pd.ExcelWriter("EC-data.xlsx", engine="openpyxl") as writer:
    data_EC.to_excel(writer, 
                     sheet_name='Time dependence', 
                     index=False)
    
    anl_EC_t.to_excel(writer,
                      sheet_name='Average', 
                      index=False)

with pd.ExcelWriter("FE-data.xlsx", engine="openpyxl") as writer:
    anl_product_mol_t.to_excel(writer, 
                               sheet_name='Productivity (mol)', 
                               index=False)
    anl_product_FE_t.to_excel(writer, 
                              sheet_name='Selectivity (%)', 
                              index=False)
    anl_product_FE.to_excel(writer, 
                            sheet_name='Average Selectivity (%)', 
                            index=False)


# In[33]:


# Create a new Word document
doc = Document()

# Create or get the footer
section = doc.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

# Add "Page " text
#paragraph.text = "Page "

# Add PAGE field
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.text = "PAGE"  # Word PAGE field
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

paragraph.alignment = 1  # center alignment

#doc.add_heading('Report of analyzed data', level=1)

# Add a paragraph
#doc.add_paragraph("This is a generated figure:")

# Add an image (from file)
doc.add_picture('Ewe.png')
doc.add_picture('Ecell.png')
doc.add_picture('Ewe.png')
doc.add_picture('FE_t.png')
doc.add_picture('FE.png')

# Save the document
doc.save('report_EC.docx')


# In[ ]:




